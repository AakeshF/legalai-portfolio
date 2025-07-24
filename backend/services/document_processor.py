# services/document_processor.py - Enhanced document processing with page counting
import os
import asyncio
import json
import re
from typing import Optional, Dict, Any
from pathlib import Path
import mimetypes
from datetime import datetime

# Text extraction imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import from root level files
from database import get_db
from models import Document

# Import AI service for metadata extraction
from services.local_ai_service import AIService

# Import new RAG components
from services.document_chunker import LegalDocumentChunker
from services.embedding_service import get_embedding_service
from models import DocumentChunk, ChunkEmbedding, EmbeddingModel


class DocumentProcessor:
    """Enhanced document processor with better text extraction and error handling"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'text/plain': '.txt'
        }
        self.ai_service = AIService()
        
        # Initialize RAG components
        self.chunker = LegalDocumentChunker(
            chunk_size=512,
            chunk_overlap=100,
            min_chunk_size=100,
            max_chunk_size=1024
        )
        self.embedding_service = get_embedding_service()
        
        print(f"📄 Document Processor initialized")
        print(f"📄 PDF support: {PDF_AVAILABLE}")
        print(f"📄 DOCX support: {DOCX_AVAILABLE}")
        print(f"🤖 AI Service integration: {'Demo Mode' if self.ai_service.demo_mode else 'Active'}")
        print(f"🔍 Semantic Search: Enabled with {self.embedding_service.model_name}")
    
    async def process_document(self, document_id: str, file_path: str, extract_metadata: bool = False, organization_id: str = None) -> bool:
        """Process a document: extract text, generate summary, update status, optionally extract metadata
        
        Args:
            document_id: Document ID to process
            file_path: Path to the document file
            extract_metadata: Whether to extract legal metadata
            organization_id: Organization ID for security verification
        """
        
        db = next(get_db())
        
        # Build query with optional organization filter for security
        query = db.query(Document).filter(Document.id == document_id)
        if organization_id:
            query = query.filter(Document.organization_id == organization_id)
        
        document = query.first()
        
        if not document:
            print(f"❌ Document {document_id} not found or unauthorized access")
            return False
        
        try:
            print(f"🔄 Processing document: {document.filename}")
            
            # Update status to processing
            document.processing_status = "processing"
            db.commit()
            print(f"📊 Status updated to 'processing' for {document.filename}")
            
            # Extract text content and get page count
            extraction_result = await self._extract_text_and_pages(file_path, document.filename)
            
            if not extraction_result or not extraction_result.get('text'):
                print(f"❌ No text extracted from {document.filename}")
                document.processing_status = "failed"
                document.error_message = "Could not extract text content"
                db.commit()
                return False
            
            extracted_text = extraction_result['text']
            page_count = extraction_result.get('page_count')
            
            # Update document with extracted content
            document.extracted_content = extracted_text
            document.page_count = page_count
            
            # Generate summary
            summary = self._generate_summary(extracted_text, document.filename)
            document.summary = summary
            
            # Extract legal metadata if requested and AI service is available
            if extract_metadata and not self.ai_service.demo_mode:
                print(f"🔍 Extracting legal metadata for {document.filename}")
                metadata = await self.extract_legal_metadata(extracted_text, document.filename)
                
                # Store metadata as JSON string in the document
                import json
                document.legal_metadata = json.dumps(metadata)
                print(f"📋 Metadata extracted: {metadata.get('document_type', 'unknown')} with {len(metadata.get('parties', []))} parties")
            elif extract_metadata and self.ai_service.demo_mode:
                print(f"⚠️ Metadata extraction skipped - AI service in demo mode")
                # Use fallback extraction in demo mode
                metadata = self._fallback_metadata_extraction(extracted_text)
                import json
                document.legal_metadata = json.dumps(metadata)
            
            # Generate chunks and embeddings for semantic search
            await self._generate_chunks_and_embeddings(
                db, document, extracted_text, metadata if extract_metadata else None
            )
            
            # Update status to completed
            document.processing_status = "completed"
            document.processed_timestamp = datetime.utcnow()
            db.commit()
            
            print(f"✅ Successfully processed {document.filename}")
            print(f"📄 Extracted {len(extracted_text)} characters")
            print(f"📄 Page count: {page_count}")
            print(f"📝 Generated summary: {summary[:100]}...")
            
            return True
            
        except Exception as e:
            print(f"❌ Error processing document {document.filename}: {str(e)}")
            document.processing_status = "failed"
            document.error_message = str(e)
            db.commit()
            return False
        finally:
            db.close()
    
    async def _extract_text_and_pages(self, file_path: str, filename: str) -> Optional[Dict[str, Any]]:
        """Extract text and page count from various file formats"""
        
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        # Get file type
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = Path(filename).suffix.lower()
        
        print(f"🔍 Processing file: {filename}")
        print(f"🔍 MIME type: {mime_type}")
        print(f"🔍 Extension: {file_extension}")
        
        try:
            # PDF files
            if file_extension == '.pdf' or mime_type == 'application/pdf':
                return await self._extract_pdf_text_and_pages(file_path)
            
            # DOCX files
            elif file_extension == '.docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_docx_text_and_pages(file_path)
            
            # TXT files
            elif file_extension == '.txt' or mime_type == 'text/plain':
                return await self._extract_txt_text_and_pages(file_path)
            
            else:
                print(f"❌ Unsupported file type: {mime_type} / {file_extension}")
                return None
                
        except Exception as e:
            print(f"❌ Text extraction error for {filename}: {str(e)}")
            return None
    
    async def _extract_pdf_text_and_pages(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Extract text and page count from PDF files with memory optimization"""
        
        if not PDF_AVAILABLE:
            print("❌ PyPDF2 not available for PDF processing")
            return {
                'text': "PDF processing not available. Please install PyPDF2.",
                'page_count': None
            }
        
        try:
            # Check file size first
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            print(f"📊 PDF file size: {file_size_mb:.2f} MB")
            
            # For large files, use streaming approach
            if file_size_mb > 10:  # Files larger than 10MB
                return await self._extract_pdf_streaming(file_path)
            
            text_content = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                print(f"📄 PDF has {page_count} pages")
                
                # Process in batches to avoid memory buildup
                batch_size = 50
                for batch_start in range(0, page_count, batch_size):
                    batch_end = min(batch_start + batch_size, page_count)
                    batch_text = []
                    
                    for page_num in range(batch_start, batch_end):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text.strip():
                                batch_text.append(f"\n--- Page {page_num + 1} ---\n")
                                batch_text.append(page_text)
                            # Clear page object to free memory
                            del page
                        except Exception as e:
                            print(f"⚠️ Error extracting page {page_num + 1}: {str(e)}")
                            continue
                    
                    # Join batch and add to content
                    if batch_text:
                        text_content.append("".join(batch_text))
                    
                    # Clear batch to free memory
                    del batch_text
            
            full_text = "\n".join(text_content)
            
            # Clear intermediate list
            del text_content
            
            if not full_text.strip():
                full_text = "PDF appears to be empty or contains only images. Consider using OCR for scanned documents."
            
            print(f"✅ Extracted {len(full_text)} characters from PDF")
            
            return {
                'text': full_text,
                'page_count': page_count
            }
            
        except Exception as e:
            print(f"❌ PDF extraction error: {str(e)}")
            return {
                'text': f"Error processing PDF: {str(e)}",
                'page_count': None
            }
    
    async def _extract_docx_text_and_pages(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Extract text from DOCX files (page count not easily available)"""
        
        if not DOCX_AVAILABLE:
            print("❌ python-docx not available for DOCX processing")
            return {
                'text': "DOCX processing not available. Please install python-docx.",
                'page_count': None
            }
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            full_text = "\n\n".join(paragraphs)
            
            if not full_text.strip():
                full_text = "DOCX file appears to be empty."
            
            print(f"✅ Extracted {len(full_text)} characters from DOCX")
            
            # Estimate page count based on text length (rough estimate)
            estimated_pages = max(1, len(full_text) // 2500) if full_text.strip() else 1
            
            return {
                'text': full_text,
                'page_count': estimated_pages
            }
            
        except Exception as e:
            print(f"❌ DOCX extraction error: {str(e)}")
            return {
                'text': f"Error processing DOCX: {str(e)}",
                'page_count': None
            }
    
    async def _extract_txt_text_and_pages(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Extract text from TXT files"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                    if content.strip():
                        print(f"✅ Successfully read TXT file with {encoding} encoding")
                        print(f"✅ Extracted {len(content)} characters from TXT")
                        
                        # Estimate page count based on line breaks or character count
                        line_count = len(content.split('\n'))
                        estimated_pages = max(1, line_count // 50)  # ~50 lines per page
                        
                        return {
                            'text': content,
                            'page_count': estimated_pages
                        }
                        
                except UnicodeDecodeError:
                    continue
            
            print("❌ Could not decode TXT file with any encoding")
            return {
                'text': "Error: Could not decode text file. Please check file encoding.",
                'page_count': None
            }
            
        except Exception as e:
            print(f"❌ TXT extraction error: {str(e)}")
            return {
                'text': f"Error processing TXT: {str(e)}",
                'page_count': None
            }
    
    def _generate_summary(self, text: str, filename: str) -> str:
        """Generate a summary of the document content"""
        
        if not text or len(text.strip()) < 50:
            return "Document appears to be empty or very short."
        
        # Basic document analysis
        word_count = len(text.split())
        char_count = len(text)
        line_count = len(text.split('\n'))
        
        # Try to determine document type based on content
        text_lower = text.lower()
        doc_type = "Document"
        
        # Legal document detection
        legal_keywords = [
            'contract', 'agreement', 'whereas', 'party', 'parties', 'hereby',
            'shall', 'terms', 'conditions', 'liability', 'breach', 'damages',
            'court', 'jurisdiction', 'clause', 'provision', 'execution'
        ]
        
        legal_score = sum(1 for keyword in legal_keywords if keyword in text_lower)
        
        if legal_score >= 3:
            doc_type = "Legal Document"
            
            # More specific legal document types
            if any(word in text_lower for word in ['lease', 'rent', 'tenant', 'landlord']):
                doc_type = "Lease Agreement"
            elif any(word in text_lower for word in ['employment', 'employee', 'employer', 'salary']):
                doc_type = "Employment Contract"
            elif any(word in text_lower for word in ['non-disclosure', 'confidential', 'nda']):
                doc_type = "Non-Disclosure Agreement"
            elif any(word in text_lower for word in ['purchase', 'sale', 'buyer', 'seller']):
                doc_type = "Purchase Agreement"
        
        # Extract first meaningful sentence as preview
        sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 20]
        preview = sentences[0][:200] + "..." if sentences else "No meaningful content preview available."
        
        summary = f"""Document Type: {doc_type}
File: {filename}
Content: {word_count} words, {line_count} lines
Preview: {preview}"""
        
        if legal_score >= 3:
            summary += f"\nLegal Keywords Found: {legal_score} (appears to be a legal document)"
        
        return summary
    
    async def extract_legal_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract structured legal metadata using AI service"""
        
        if not text or len(text.strip()) < 50:
            return {
                "error": "Insufficient text for metadata extraction",
                "parties": [],
                "dates": [],
                "monetary_amounts": [],
                "document_type": "unknown"
            }
        
        # Prepare the combined prompt for efficient API usage
        metadata_prompt = """Analyze this legal document and extract the following information in JSON format:

{
  "document_type": "Type of legal document (contract, lease, motion, complaint, etc.)",
  "parties": [
    {"role": "role name", "name": "party name"}
  ],
  "dates": [
    {"event": "event description", "date": "MM/DD/YYYY"}
  ],
  "monetary_amounts": [
    {"description": "what the amount is for", "amount": "$X,XXX.XX"}
  ],
  "key_obligations": ["list of main obligations or requirements"],
  "jurisdiction": "applicable jurisdiction if mentioned",
  "governing_law": "governing law if mentioned"
}

Focus on accuracy. Only include information explicitly stated in the document.
If information is not found, use empty arrays or "not specified".

Document content:
"""
        
        try:
            # Use the AI service to extract metadata
            response = await self.ai_service._call_ai_api([
                {"role": "system", "content": "You are a legal document analyzer. Extract information and return valid JSON only."},
                {"role": "user", "content": metadata_prompt + text[:8000]}  # Limit text to avoid token limits
            ])
            
            # Try to parse the JSON response
            import json
            
            # Find JSON in the response (AI might include explanation text)
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_str = response[json_start:json_end]
                metadata = json.loads(json_str)
                
                # Add extraction timestamp
                metadata['extraction_timestamp'] = datetime.utcnow().isoformat()
                metadata['extraction_method'] = 'ai_provider'
                
                print(f"✅ Successfully extracted metadata for {filename}")
                return metadata
            else:
                raise ValueError("No JSON found in response")
                
        except Exception as e:
            print(f"❌ Error extracting metadata: {str(e)}")
            
            # Fallback to basic extraction using regex patterns
            return self._fallback_metadata_extraction(text)
    
    def _fallback_metadata_extraction(self, text: str) -> Dict[str, Any]:
        """Fallback metadata extraction using patterns"""
        
        import re
        
        # Basic patterns for common legal document elements
        metadata = {
            "document_type": "unknown",
            "parties": [],
            "dates": [],
            "monetary_amounts": [],
            "key_obligations": [],
            "jurisdiction": "not specified",
            "governing_law": "not specified",
            "extraction_method": "pattern_matching"
        }
        
        # Try to determine document type
        text_lower = text.lower()
        if 'lease' in text_lower and 'tenant' in text_lower:
            metadata['document_type'] = 'lease'
        elif 'purchase' in text_lower and 'agreement' in text_lower:
            metadata['document_type'] = 'purchase_agreement'
        elif 'employment' in text_lower and 'agreement' in text_lower:
            metadata['document_type'] = 'employment_contract'
        elif 'complaint' in text_lower and 'defendant' in text_lower:
            metadata['document_type'] = 'complaint'
        
        # Extract dates
        date_pattern = r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
        dates = re.findall(date_pattern, text)
        metadata['dates'] = [{"event": "mentioned", "date": date} for date in dates[:5]]
        
        # Extract monetary amounts
        money_pattern = r'\$[\d,]+(?:\.\d{2})?'
        amounts = re.findall(money_pattern, text)
        metadata['monetary_amounts'] = [{"description": "mentioned", "amount": amt} for amt in amounts[:5]]
        
        return metadata


    async def _extract_pdf_streaming(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Extract text from large PDFs using streaming approach"""
        print("📚 Using streaming extraction for large PDF")
        
        try:
            page_count = 0
            total_chars = 0
            max_chars = 5_000_000  # 5MB text limit
            
            # First pass: get page count
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
            
            print(f"📄 Large PDF has {page_count} pages")
            
            # Second pass: extract text in chunks
            extracted_chunks = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(page_count):
                    if total_chars >= max_chars:
                        extracted_chunks.append(f"\n\n[... Truncated at page {page_num + 1} of {page_count} due to size limit ...]")
                        break
                    
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        
                        if page_text.strip():
                            chunk = f"\n--- Page {page_num + 1} ---\n{page_text}"
                            chunk_size = len(chunk)
                            
                            if total_chars + chunk_size <= max_chars:
                                extracted_chunks.append(chunk)
                                total_chars += chunk_size
                            else:
                                # Partial page to fit within limit
                                remaining = max_chars - total_chars
                                extracted_chunks.append(chunk[:remaining] + "\n[... Truncated ...]")
                                total_chars = max_chars
                                break
                        
                        # Explicitly delete page to free memory
                        del page
                        
                    except Exception as e:
                        print(f"⚠️ Error extracting page {page_num + 1}: {str(e)}")
                        continue
                    
                    # Yield control periodically
                    if page_num % 10 == 0:
                        await asyncio.sleep(0)  # Allow other tasks to run
            
            full_text = "".join(extracted_chunks)
            del extracted_chunks  # Free memory
            
            print(f"✅ Extracted {len(full_text)} characters from large PDF (capped at {max_chars})")
            
            return {
                'text': full_text,
                'page_count': page_count
            }
            
        except Exception as e:
            print(f"❌ Streaming PDF extraction error: {str(e)}")
            return {
                'text': f"Error processing large PDF: {str(e)}",
                'page_count': None
            }
    
    async def _generate_chunks_and_embeddings(
        self,
        db,
        document: Document,
        text: str,
        metadata: Optional[Dict] = None
    ):
        """Generate chunks and embeddings for semantic search"""
        try:
            print(f"🔍 Generating chunks for {document.filename}")
            
            # Check if embedding model exists, create if not
            embedding_model = db.query(EmbeddingModel).filter(
                EmbeddingModel.name == self.embedding_service.model_name
            ).first()
            
            if not embedding_model:
                embedding_model = EmbeddingModel(
                    name=self.embedding_service.model_name,
                    dimension=self.embedding_service.dimension,
                    provider='local',
                    description='Local sentence transformer model'
                )
                db.add(embedding_model)
                db.commit()
            
            # Generate chunks
            chunks = self.chunker.chunk_document(
                text=text,
                document_type=metadata.get('document_type', 'general') if metadata else 'general',
                metadata=metadata
            )
            
            print(f"📄 Generated {len(chunks)} chunks")
            
            # Store chunks in database
            chunk_objects = []
            for chunk in chunks:
                chunk_obj = DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk.chunk_index,
                    content=chunk.content,
                    tokens=chunk.tokens,
                    start_char=chunk.start_char,
                    end_char=chunk.end_char,
                    chunk_metadata=chunk.metadata
                )
                db.add(chunk_obj)
                chunk_objects.append(chunk_obj)
            
            # Commit chunks to get IDs
            db.commit()
            
            # Generate embeddings in batches
            print(f"🧮 Generating embeddings for {len(chunk_objects)} chunks")
            
            chunk_texts = [chunk.content for chunk in chunk_objects]
            embedding_results = await self.embedding_service.generate_embeddings_batch(
                texts=chunk_texts,
                metadata=[{'chunk_id': chunk.id} for chunk in chunk_objects]
            )
            
            # Store embeddings
            for chunk_obj, embedding_result in zip(chunk_objects, embedding_results):
                embedding = ChunkEmbedding(
                    chunk_id=chunk_obj.id,
                    model_id=embedding_model.id,
                    embedding=json.dumps(embedding_result.embedding),  # Convert to JSON for SQLite
                    encoding_time_ms=embedding_result.encoding_time_ms
                )
                db.add(embedding)
                
                # Mark chunk as having embedding
                chunk_obj.embedding_generated = True
            
            # Update document metadata
            document.chunks_generated = True
            document.embeddings_generated = True
            document.embedding_model_id = embedding_model.id
            document.chunk_count = len(chunks)
            document.last_embedded_at = datetime.utcnow()
            
            db.commit()
            
            print(f"✅ Successfully generated {len(chunks)} chunks with embeddings")
            
        except Exception as e:
            print(f"❌ Error generating chunks/embeddings: {str(e)}")
            # Don't fail the whole document processing
            db.rollback()


# Global processor instance
document_processor = DocumentProcessor()